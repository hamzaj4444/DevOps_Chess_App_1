from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── Helper utilities ────────────────────────────────────────────────────────

def set_heading(text, level, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_paragraph(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_code_block(text):
    for line in text.strip().splitlines():
        p = doc.add_paragraph(line)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

def shade_table_cell(cell, hex_color="D0E4F7"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, header_color="1565C0", header_text_color=RGBColor(255, 255, 255)):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = header_text_color
        shade_table_cell(hdr_cells[i], header_color)
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val
    return table

# ─── Page Setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ─── COVER PAGE ──────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("DevOps Chess Application")
title_run.bold = True
title_run.font.size = Pt(30)
title_run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run("Full Project Technical Report")
sub_run.font.size = Pt(18)
sub_run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc.add_run(
    "A comprehensive report documenting the architecture, design decisions,\n"
    "implementation details, DevOps practices, and Kubernetes migration\n"
    "of the AI-powered Chess application."
)
desc_run.font.size = Pt(11)
desc_run.italic = True
desc_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f"Report Date: {datetime.date.today().strftime('%B %d, %Y')}")
date_run.font.size = Pt(12)
date_run.bold = True

doc.add_page_break()

# ─── 1. EXECUTIVE SUMMARY ────────────────────────────────────────────────────
set_heading("1. Executive Summary", 1, color=(0x15, 0x65, 0xC0))

add_paragraph(
    "The DevOps Chess Application is a full-stack, enterprise-grade AI-powered chess platform "
    "built from the ground up using modern software engineering and DevOps best practices. "
    "The application features an intelligent AI opponent powered by a PyTorch neural network, "
    "a modern Glassmorphism web frontend, a secure FastAPI backend, and a complete observability stack "
    "comprising Prometheus, Grafana, Loki, and Promtail."
)
doc.add_paragraph()
add_paragraph(
    "The project demonstrates the full lifecycle of a production-ready software product: "
    "from local development and containerization, through continuous integration and automated "
    "security scanning, to full orchestration with Kubernetes on a local Minikube cluster."
)
doc.add_paragraph()

add_paragraph("Key Highlights:", bold=True)
highlights = [
    "AI Chess engine powered by a custom PyTorch neural network running on CPU.",
    "Secure and authenticated REST API built with FastAPI and JWT tokens.",
    "Modern Glassmorphism single-page UI with real-time chess board interaction.",
    "Full observability stack: Prometheus metrics, Loki + Promtail log aggregation, Grafana dashboards.",
    "GitHub Actions CI/CD pipeline with automated testing, Docker Hub publishing, and Trivy security scanning.",
    "Complete Kubernetes migration: all 6 services deployed to a local Minikube cluster.",
    "Auto-provisioned Grafana datasources and dashboards via Kubernetes ConfigMaps.",
]
for h in highlights:
    add_bullet(h)

doc.add_page_break()

# ─── 2. PROJECT ARCHITECTURE ─────────────────────────────────────────────────
set_heading("2. Overall Architecture", 1, color=(0x15, 0x65, 0xC0))
add_paragraph(
    "The application is organized into the following core layers, each independently containerized "
    "and orchestrated via Kubernetes. The architecture prioritizes separation of concerns, observability, "
    "and security at each layer."
)
doc.add_paragraph()
add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Reverse Proxy",    "Nginx (Alpine)",           "HTTP request routing, load balancing, static content caching"],
        ["Backend API",      "Python 3.12 + FastAPI",    "Business logic, authentication, chess engine interface, metrics"],
        ["AI Chess Engine",  "PyTorch + python-chess",   "Neural network move generation, game state management"],
        ["Metrics",          "Prometheus Client",        "Exporting custom counters, histograms, gauges at /metrics"],
        ["Log Shipping",     "Grafana Promtail",         "Scraping container logs via Docker socket and forwarding to Loki"],
        ["Log Aggregation",  "Grafana Loki",             "High-scale log storage and label-based querying"],
        ["Visualization",    "Grafana",                  "Dashboards for real-time metrics and logs"],
        ["Orchestration",    "Kubernetes (Minikube)",    "Pod scheduling, service discovery, persistent storage, health management"],
    ]
)

doc.add_page_break()

# ─── 3. BACKEND ──────────────────────────────────────────────────────────────
set_heading("3. Backend Development — FastAPI", 1, color=(0x15, 0x65, 0xC0))

set_heading("3.1 Application Structure", 2)
add_paragraph(
    "The backend is organized as a clean Python package under the `app/` directory, with strict "
    "separation between authentication, business logic, metrics, and routing:"
)
add_code_block(
    "app/\n"
    "├── main.py          # FastAPI app entry point, middleware, top-level routes\n"
    "├── auth.py          # JWT token creation and validation\n"
    "├── users.py         # User lookup and password hashing\n"
    "├── metrics.py       # Prometheus counters, histograms, gauges\n"
    "├── stats.py         # Chess.com public API integration\n"
    "├── logger.py        # Centralized structured logger\n"
    "├── chess/\n"
    "│   └── routes.py    # Chess game lifecycle API routes\n"
    "└── static/chess/    # Frontend assets (HTML, CSS, JS)"
)
doc.add_paragraph()

set_heading("3.2 API Endpoints", 2)
add_table(
    ["Method", "Endpoint", "Auth Required", "Description"],
    [
        ["GET",  "/health",         "No",  "Application liveness health check"],
        ["POST", "/login",          "No",  "Authenticate user, receive JWT token"],
        ["GET",  "/stats",          "Yes", "Retrieve Chess.com player statistics"],
        ["GET",  "/metrics",        "No",  "Prometheus-formatted metrics endpoint"],
        ["POST", "/chess/game/new", "No",  "Initialize a new chess game session"],
        ["POST", "/chess/game/{id}","No",  "Submit a player move, receive bot response"],
        ["GET",  "/chess/move",     "No",  "Stateless: get single best bot move for a FEN"],
    ]
)
doc.add_paragraph()

set_heading("3.3 Authentication System", 2)
add_paragraph(
    "The application uses JSON Web Tokens (JWT) for securing protected endpoints. "
    "Passwords are hashed using bcrypt via the `passlib` library. The current implementation uses "
    "an in-memory user dictionary, which is appropriate for the current development phase but is "
    "identified as a future improvement candidate to replace with a persistent database."
)
doc.add_paragraph()

set_heading("3.4 Custom Metrics", 2)
add_paragraph("The following Prometheus metrics are exported from `app/metrics.py`:")
add_table(
    ["Metric Name", "Type", "Description"],
    [
        ["http_requests_total",            "Counter",   "Total HTTP requests, labelled by method, endpoint, status"],
        ["http_request_latency_seconds",   "Histogram", "Request latency per endpoint"],
        ["login_failures_total",           "Counter",   "Total failed login attempts"],
        ["chess_api_errors_total",         "Counter",   "Errors fetching data from Chess.com API"],
        ["chess_games_total",              "Counter",   "Total chess games started"],
        ["chess_moves_total",              "Counter",   "Total moves made, labeled by color (player/bot)"],
        ["chess_bot_move_duration_seconds","Histogram", "AI computation time per move"],
        ["chess_games_active",             "Gauge",     "Current number of concurrently active games"],
        ["chess_game_outcomes_total",      "Counter",   "Game results labeled by outcome type"],
    ]
)

doc.add_page_break()

# ─── 4. AI CHESS ENGINE ──────────────────────────────────────────────────────
set_heading("4. AI Chess Engine", 1, color=(0x15, 0x65, 0xC0))
add_paragraph(
    "The chess AI is powered by a custom PyTorch neural network model. This represents one of "
    "the most distinctive engineering aspects of the project — the AI bot is not a rules-based "
    "engine (such as Stockfish), but instead uses a learned model to evaluate board positions and "
    "select moves."
)
doc.add_paragraph()
add_paragraph("Game Lifecycle Management:", bold=True)
add_bullet("A unique game session ID is generated upon calling `POST /chess/game/new`.")
add_bullet("Game state is maintained in memory across multiple `POST /chess/game/{id}` requests.")
add_bullet("Both stateful (session-based) and stateless (single FEN query) modes are supported.")
add_bullet("The bot move latency is measured and exported as a Prometheus histogram metric.")
doc.add_paragraph()
add_paragraph("Dependencies:", bold=True)
add_bullet("torch (CPU-only build via PyTorch's dedicated index URL for smaller image size)")
add_bullet("python-chess — FEN parsing, legal move generation, move application")
add_bullet("numpy — numerical support for model input tensor construction")

doc.add_page_break()

# ─── 5. FRONTEND ─────────────────────────────────────────────────────────────
set_heading("5. Frontend Engineering", 1, color=(0x15, 0x65, 0xC0))
add_paragraph(
    "The user interface is a single-page web application served as static files by FastAPI "
    "at the `/game/` path. It is built with vanilla HTML, CSS, and JavaScript, deliberately "
    "avoiding heavy frameworks to keep the frontend lightweight within the Docker image."
)
doc.add_paragraph()
set_heading("5.1 Design System — Glassmorphism", 2)
add_paragraph(
    "The UI adopts a Glassmorphism aesthetic, characterized by frosted-glass panels, "
    "semi-transparent backgrounds, blurred backdrops, and vibrant accent colors. This gives the "
    "application a modern, premium visual appearance that stands out from typical admin interfaces."
)
doc.add_paragraph()
set_heading("5.2 Chess Board Interaction", 2)
add_paragraph(
    "The `chess.js` library is integrated to handle all client-side board logic, including legal "
    "move validation, check detection, and checkmate recognition. Player moves are submitted "
    "asynchronously via `fetch()` to the FastAPI backend, which returns the AI-chosen move and "
    "updates the rendered board accordingly."
)

doc.add_page_break()

# ─── 6. CONTAINERISATION ─────────────────────────────────────────────────────
set_heading("6. Containerisation with Docker", 1, color=(0x15, 0x65, 0xC0))

set_heading("6.1 Multi-Stage Dockerfile", 2)
add_paragraph(
    "A multi-stage build strategy is used to produce a lean, secure production image. "
    "This approach keeps all build-time tools (pip, wheel files, source) out of the final "
    "runtime image, reducing both image size and attack surface."
)
doc.add_paragraph()
add_paragraph("Stage 1 — Builder:", bold=True)
add_bullet("Base: `python:3.12-slim-bookworm`")
add_bullet("Installs all PyPI dependencies (including the large PyTorch CPU wheel) via pip into /install")
add_bullet("Does NOT include application source code")
doc.add_paragraph()
add_paragraph("Stage 2 — Runtime:", bold=True)
add_bullet("Base: fresh `python:3.12-slim-bookworm` (no build tools)")
add_bullet("Applies latest OS package security updates (`apt-get upgrade`)")
add_bullet("Creates a non-root `appuser` for process isolation")
add_bullet("Copies only compiled dependencies from Stage 1 and the `app/` source")
add_bullet("Runs as non-root user on port 8000 via Uvicorn ASGI server")
doc.add_paragraph()

set_heading("6.2 Docker Compose Stack", 2)
add_paragraph(
    "The `docker-compose.yml` defines a complete local development stack with 6 services "
    "all networked on a shared `backend` Docker bridge network:"
)
add_table(
    ["Service", "Image", "Port", "Role"],
    [
        ["app",       "hamzaj4444/devops-chess-app (built)", "8000", "FastAPI backend + AI engine"],
        ["nginx",     "nginx:alpine",     "80",   "Reverse proxy routing to the backend"],
        ["prometheus","prom/prometheus",  "9090", "Metrics scraper and storage"],
        ["loki",      "grafana/loki",     "3100", "Log aggregation and query engine"],
        ["promtail",  "grafana/promtail", "9080", "Log collector from Docker socket"],
        ["grafana",   "grafana/grafana",  "3000", "Dashboard and visualization UI"],
    ]
)

doc.add_page_break()

# ─── 7. OBSERVABILITY ────────────────────────────────────────────────────────
set_heading("7. Full Observability Stack", 1, color=(0x15, 0x65, 0xC0))
add_paragraph(
    "The project implements the three pillars of observability: Metrics, Logging, and "
    "Visualization — wired together into a production-grade monitoring pipeline."
)
doc.add_paragraph()

set_heading("7.1 Prometheus — Metrics", 2)
add_paragraph(
    "Prometheus is configured to scrape the FastAPI `/metrics` endpoint every 15 seconds. "
    "The `prometheus.yml` configures a single scrape job targeting `chess-app:8000` (resolved "
    "via internal service DNS in both Docker Compose and Kubernetes)."
)
doc.add_paragraph()

set_heading("7.2 Loki + Promtail — Logs", 2)
add_paragraph(
    "Promtail is configured as a DaemonSet (in Kubernetes) to read container logs directly "
    "from the Docker socket at `/var/run/docker.sock`. It applies relabeling rules to extract "
    "clean container names from Kubernetes-style pod names (e.g., `k8s_chess-app_chess-app-785c...`) "
    "using the regex pattern `/k8s_([^_]+)_.*`. Logs are forwarded in real-time to Loki, "
    "which indexes them using stream labels for fast querying."
)
doc.add_paragraph()

set_heading("7.3 Grafana — Dashboards", 2)
add_paragraph(
    "Grafana connects to both Prometheus and Loki as datasources. The `chess-dashboard.json` "
    "file defines a pre-built dashboard with the following panels:"
)
add_table(
    ["Panel", "Type", "Data Source", "Query"],
    [
        ["Games Started Rate",  "Time Series", "Prometheus", "rate(chess_games_total[5m]) * 60"],
        ["Active Games",        "Gauge",       "Prometheus", "chess_games_active"],
        ["Bot Move Latency P95","Time Series", "Prometheus", "histogram_quantile(0.95, rate(chess_bot_move_duration_seconds_bucket[5m]))"],
        ["Game Outcomes",       "Pie Chart",   "Prometheus", "sum by (result)(chess_game_outcomes_total)"],
        ["Chess App Logs",      "Logs Panel",  "Loki",       '{container="chess-app"} |= "chess"'],
    ]
)
doc.add_paragraph()
add_paragraph(
    "In Kubernetes, the dashboard and both datasources are auto-provisioned via ConfigMaps injected "
    "into Grafana's provisioning directory at startup. This ensures dashboards are never lost when "
    "the pod is rescheduled or the PVC is recreated."
)

doc.add_page_break()

# ─── 8. CI/CD PIPELINE ───────────────────────────────────────────────────────
set_heading("8. CI/CD Pipeline — GitHub Actions", 1, color=(0x15, 0x65, 0xC0))
add_paragraph(
    "The `.github/workflows/ci.yml` pipeline is a four-stage automated delivery pipeline "
    "that runs on every push or pull request to the `main` branch. Each stage is a dependent job "
    "that only executes if the previous stage passes."
)
doc.add_paragraph()
add_table(
    ["Stage", "Job Name", "Trigger", "Key Actions"],
    [
        ["1 — Test",     "test",            "Push & PR to main", "Install deps, run pytest with PYTHONPATH set"],
        ["2 — Build",    "build-and-push",  "Push to main only", "Docker Buildx multi-arch build, push to Docker Hub with :latest and :sha tags. Uses GHA layer cache."],
        ["3 — Security", "security-scan",   "After build",       "Trivy scan for CRITICAL/HIGH CVEs. Uploads SARIF report as artifact. Fails pipeline on unfixed CVEs."],
        ["4 — Deploy",   "deploy",          "After scan passes", "Triggers Render webhook deploy hook via curl POST."],
    ]
)
doc.add_paragraph()
add_paragraph("Required GitHub Secrets:", bold=True)
add_table(
    ["Secret Name", "Purpose"],
    [
        ["DOCKERHUB_USERNAME",    "Docker Hub account username for image tagging and pushing"],
        ["DOCKER_PASSWORD",       "Docker Hub access token for authentication during push"],
        ["RENDER_DEPLOY_HOOK_URL","Webhook URL to trigger re-deployment on Render cloud platform"],
    ]
)

doc.add_page_break()

# ─── 9. KUBERNETES MIGRATION ─────────────────────────────────────────────────
set_heading("9. Kubernetes Migration", 1, color=(0x15, 0x65, 0xC0))
add_paragraph(
    "The entire application stack was migrated from Docker Compose to a local Kubernetes cluster "
    "managed by Minikube. This involved translating every service from Docker Compose definitions "
    "into proper Kubernetes resource manifests."
)
doc.add_paragraph()

set_heading("9.1 Cluster Setup", 2)
add_paragraph("The following tools were installed to enable the Kubernetes environment:")
add_bullet("kubectl — Kubernetes CLI (v1.35+)")
add_bullet("minikube — Local Kubernetes cluster runner using Docker as the hypervisor driver")
add_paragraph("Minikube is started using the Docker driver:", italic=True)
add_code_block("minikube start --driver=docker")
doc.add_paragraph()

set_heading("9.2 Kubernetes Manifests Inventory", 2)
add_table(
    ["File", "Resources", "Service"],
    [
        ["chess-deployment.yaml",         "Deployment, Service, Secret",         "FastAPI chess backend"],
        ["nginx-deployment.yaml",         "Deployment",                          "Nginx reverse proxy"],
        ["nginx-service.yaml",            "Service (ClusterIP)",                 "Nginx service discovery"],
        ["nginx-configmap.yaml",          "ConfigMap",                           "Nginx virtual host config"],
        ["prometheus.yaml",               "Deployment, Service, ConfigMap, PVC", "Prometheus metrics server"],
        ["loki.yaml",                     "Deployment, Service, ConfigMap, PVC", "Loki log aggregation"],
        ["promtail.yaml",                 "DaemonSet, ConfigMap",                "Promtail log collection"],
        ["grafana.yaml",                  "Deployment, Service, PVC",            "Grafana visualization"],
        ["grafana-provisioning-cm.yaml",  "ConfigMap",                           "Grafana dashboard provider config"],
        ["grafana-datasources-cm.yaml",   "ConfigMap",                           "Auto-provisioned Prometheus + Loki datasources"],
        ["grafana-dashboard-cm.yaml",     "ConfigMap",                           "Auto-provisioned Chess App dashboard JSON"],
    ]
)
doc.add_paragraph()

set_heading("9.3 Key Design Decisions", 2)
add_paragraph("Building the chess-app image inside Minikube:", bold=True)
add_paragraph(
    "Because Minikube uses an isolated Docker daemon, the application image must be built "
    "within Minikube's environment rather than the host's. This is achieved by pointing the "
    "Docker CLI to Minikube's daemon before building:"
)
add_code_block(
    "eval $(minikube docker-env)\n"
    "docker build -t hamzaj4444/devops-chess-app:latest .\n"
    "kubectl apply -f kubernetes/"
)
doc.add_paragraph()
add_paragraph("imagePullPolicy: IfNotPresent:", bold=True)
add_paragraph(
    "The chess-app Deployment is configured with `imagePullPolicy: IfNotPresent` to prevent "
    "Kubernetes from attempting to pull the image from a remote registry (Docker Hub) when it "
    "has already been built locally within the cluster's Docker environment."
)
doc.add_paragraph()
add_paragraph("PersistentVolumeClaims (PVCs):", bold=True)
add_paragraph(
    "Prometheus, Loki, and Grafana each use a PVC backed by Minikube's default `hostpath` "
    "dynamic provisioner. This ensures that metric data, log indices, and Grafana configurations "
    "are persisted across pod restarts."
)
doc.add_paragraph()

set_heading("9.4 Accessing Services via Minikube", 2)
add_paragraph(
    "Since Minikube runs inside a Docker container on Linux, Kubernetes ClusterIP services are "
    "not directly reachable from the host machine. Minikube provides tunneling via the following commands:"
)
add_table(
    ["Service", "Command", "Default Path"],
    [
        ["Chess App (Nginx)", "minikube service nginx-service --url",   "/game/"],
        ["Grafana",           "minikube service grafana-service --url", "/ (login: admin/admin)"],
        ["Prometheus",        "minikube service prometheus-service --url", "/graph"],
    ]
)

doc.add_page_break()

# ─── 10. VERIFIED DEPLOYMENT STATUS ──────────────────────────────────────────
set_heading("10. Verified Deployment Status", 1, color=(0x15, 0x65, 0xC0))
add_paragraph(
    "All Kubernetes pods were verified as fully running and healthy following the migration. "
    "The following table reflects the final verified state of the cluster:"
)
doc.add_paragraph()
add_table(
    ["Pod Name", "Ready", "Status", "Service"],
    [
        ["chess-app-785c45b79b-zczct",    "1/1", "Running", "FastAPI backend + AI engine"],
        ["nginx-proxy-fc9dfcfbf-7zh8h",   "1/1", "Running", "Nginx reverse proxy"],
        ["grafana-57d7b58787-htt7r",      "1/1", "Running", "Grafana dashboards"],
        ["loki-f795555ff-qxlt6",          "1/1", "Running", "Log aggregation"],
        ["prometheus-84bfdfc9bb-ff4rr",   "1/1", "Running", "Metrics collection"],
        ["promtail-hrk2n",                "1/1", "Running", "Log shipping DaemonSet"],
    ]
)
doc.add_paragraph()
add_paragraph(
    "Prometheus was additionally verified by accessing the Target Health page, which showed "
    "the chess-app scrape target as UP with an 8ms scrape duration — confirming end-to-end "
    "metrics pipeline functionality within the cluster."
)

doc.add_page_break()

# ─── 11. ISSUES ENCOUNTERED ──────────────────────────────────────────────────
set_heading("11. Issues Encountered & Resolutions", 1, color=(0x15, 0x65, 0xC0))
add_table(
    ["Issue", "Root Cause", "Resolution"],
    [
        [
            "Minikube failed to start (PROVIDER_DOCKER_VERSION_EXIT_1)",
            "Docker daemon was not running. Docker Desktop service was not found.",
            "Installed Docker Engine properly, set the Docker context to 'default', then restarted Minikube.",
        ],
        [
            "chess-app pod in ErrImagePull",
            "Minikube's internal Docker registry had no knowledge of the locally built image.",
            "Ran `eval $(minikube docker-env)` to re-point Docker CLI to Minikube, then re-built the image inside the cluster.",
        ],
        [
            "Loki, Prometheus, Promtail in CrashLoopBackOff",
            "Kubernetes `command:` field overrides the container's ENTRYPOINT, causing the runtime to try to exec the config flag string as a binary.",
            "Changed all three deployments from `command:` to `args:` so the flags are passed as arguments to the existing entrypoint.",
        ],
        [
            "Grafana Down — 'Datasource provisioning error: data source not found'",
            "Adding `uid` fields to the provisioning ConfigMap conflicted with an already-initialized SQLite database from the existing PVC.",
            "Deleted the Grafana Deployment and PVC, then re-applied with a clean slate, allowing the provisioner to write fresh uid entries.",
        ],
        [
            "'Chess App Logs' panel showed No Data in Grafana",
            "Promtail extracted the raw Kubernetes container name (`k8s_chess-app_chess-app-785c...`) which didn't match the dashboard label filter `{container=\"chess-app\"}`.",
            "Updated the Promtail relabeling regex from `/(.*)`  to `/k8s_([^_]+)_.*` to extract only the clean container name.",
        ],
    ]
)

doc.add_page_break()

# ─── 12. FUTURE IMPROVEMENTS ─────────────────────────────────────────────────
set_heading("12. Future Improvements & Roadmap", 1, color=(0x15, 0x65, 0xC0))
add_paragraph("The following enhancements are recommended for future development phases:")
doc.add_paragraph()

improvements = [
    ("Database Integration", "Replace the hardcoded `users.py` dictionary with a PostgreSQL or SQLite persistent database, using SQLAlchemy for ORM. This would support multi-user registration, profile management, and game history storage."),
    ("Dynamic Chess.com Stats", "The `chess_api_url` in `stats.py` currently hardcodes the username `dextre4`. This should be made dynamic by accepting a username parameter from the authenticated user's JWT claims."),
    ("Helm Chart Packaging", "Package all Kubernetes YAML manifests into a Helm chart, enabling versioned releases, easy environment overrides (dev/staging/production), and simplified rollbacks."),
    ("Horizontal Pod Autoscaler (HPA)", "Add an HPA for the chess-app Deployment to automatically scale replicas based on CPU usage, ensuring the AI engine remains responsive under concurrent game load."),
    ("Cloudflare Tunnel Exposure", "Use `cloudflared` to create a zero-configuration, secure public tunnel to the Nginx service, making the application accessible from the internet without port forwarding or static IPs."),
    ("Persistent Game History", "Store completed game records (moves, outcome, duration) in a database to enable leaderboards, player statistics, and replay functionality."),
    ("HTTPS / TLS Termination", "Add a cert-manager and Ingress controller to handle TLS certificate provisioning (via Let's Encrypt) for secure HTTPS access to the application."),
    ("Alerting Rules", "Define Prometheus alerting rules and connect them to Grafana Alertmanager and a notification channel (e.g., email or Slack) for production incident response."),
]
for title_text, detail in improvements:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(detail)

doc.add_page_break()

# ─── 13. CONCLUSION ──────────────────────────────────────────────────────────
set_heading("13. Conclusion", 1, color=(0x15, 0x65, 0xC0))
add_paragraph(
    "The DevOps Chess Application stands as a complete, production-grade demonstration of "
    "modern software engineering and DevOps engineering. It showcases the ability to design, "
    "build, secure, monitor, and orchestrate a non-trivial AI application using an entirely "
    "open-source toolchain."
)
doc.add_paragraph()
add_paragraph(
    "Starting from an initial FastAPI skeleton, the project was evolved through multiple "
    "engineering phases: adding an AI chess engine, crafting a polished frontend, constructing "
    "a complete observability pipeline, automating the delivery lifecycle with CI/CD, and finally "
    "achieving full Kubernetes orchestration on a local cluster — all running on a single developer machine."
)
doc.add_paragraph()
add_paragraph(
    "The project demonstrates real-world problem solving: adapting configurations between Docker "
    "Compose and Kubernetes, debugging container crashes from log output, fixing service discovery "
    "via Kubernetes DNS, and leveraging infrastructure-as-code to make dashboards and datasources "
    "fully automated and reproducible."
)
doc.add_paragraph()
add_paragraph(
    "The resulting system is a solid foundation for further growth, with a clear roadmap toward "
    "multi-user support, public cloud exposure, and Helm-based release management.",
    bold=False
)

# ─── Save ─────────────────────────────────────────────────────────────────────
output = "/home/bale/Desktop/DevOps_chess_app/DevOps_Chess_App_Report.docx"
doc.save(output)
print(f"Report saved to: {output}")
